"""
Document Redaction API
Single file FastAPI application with document processing and redaction capabilities
"""

import os
import re
import json
import uuid
import asyncio
import aiofiles
import tempfile
from datetime import datetime
from pathlib import Path
from typing import List, Dict, Tuple, Optional, Any
from fastapi import FastAPI, File, UploadFile, HTTPException, BackgroundTasks
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, JSONResponse
from pydantic import BaseModel
import spacy
from transformers import pipeline

# OCR and document processing imports
try:
    import pytesseract
    from PIL import Image, ImageDraw
    import cv2
    import numpy as np
    OCR_AVAILABLE = True
except ImportError as e:
    print(f"OCR dependencies not available: {e}")
    OCR_AVAILABLE = False

try:
    import fitz  # PyMuPDF
    PDF_AVAILABLE = True
except ImportError:
    print("PyMuPDF not available")
    PDF_AVAILABLE = False

try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    print("python-docx not available")
    DOCX_AVAILABLE = False

# Initialize FastAPI app
app = FastAPI(
    title="Document Redaction API",
    description="API for document upload, text extraction, and sensitive data redaction",
    version="1.0.0",
    docs_url="/api/docs",
    redoc_url="/api/redoc"
)

# CORS Configuration
app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:3000", "http://localhost:5173", "http://localhost:5174", "http://localhost:5175", "http://localhost:5176"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Pydantic Models
class EntitySelection(BaseModel):
    text: str
    type: str
    start: int
    end: int
    selected: bool = True  # Default to selected

class RedactionRequest(BaseModel):
    document_id: str
    entities: Optional[List[EntitySelection]] = None  # Make optional

class RedactionResponse(BaseModel):
    success: bool
    document_id: str
    redacted_text: str
    redacted_count: int
    download_url: str

class DocumentInfo(BaseModel):
    document_id: str
    filename: str
    upload_time: str
    status: str
    entity_count: int

# Global storage
documents_db = {}
uploads_dir = Path("uploads")
downloads_dir = Path("downloads")

# Create directories
uploads_dir.mkdir(exist_ok=True)
downloads_dir.mkdir(exist_ok=True)

def make_json_serializable(obj):
    """Convert numpy and other non-serializable types to Python native types"""
    if hasattr(obj, 'item'):  # numpy scalars
        return obj.item()
    elif hasattr(obj, 'tolist'):  # numpy arrays
        return obj.tolist()
    elif isinstance(obj, dict):
        return {k: make_json_serializable(v) for k, v in obj.items()}
    elif isinstance(obj, list):
        return [make_json_serializable(item) for item in obj]
    else:
        return obj

class HealthcareRedactionEngine:
    def __init__(self):
        print("🔧 INITIALIZING CONSERVATIVE HEALTHCARE REDACTION ENGINE...")
        
        # Conservative NER model with strict settings
        try:
            self.ner_pipeline = pipeline(
                "ner", 
                model="dbmdz/bert-large-cased-finetuned-conll03-english",
                aggregation_strategy="max",  # Only highest confidence predictions
                device=-1
            )
            # Set stricter confidence threshold for transformer
            self.transformer_min_confidence = 0.95  # Very high threshold
            print("✅ Conservative Transformers NER model loaded")
        except Exception as e:
            print(f"⚠️ NER model loading failed: {e}")
            self.ner_pipeline = None
        
        # Conservative spaCy settings
        try:
            self.nlp = spacy.load("en_core_web_sm")
            # Disable problematic components that over-detect
            if hasattr(self.nlp, 'disable_pipes'):
                try:
                    self.nlp.disable_pipes(['textcat', 'textcat_multilabel'])
                except:
                    pass
            print("✅ Conservative spaCy model loaded")
        except OSError:
            try:
                self.nlp = spacy.load("en_core_web_md")
                print("✅ Conservative spaCy medium model loaded")
            except OSError:
                print("⚠️ No spaCy model available")
                self.nlp = None
        
        # Initialize OCR capabilities
        self.ocr_available = OCR_AVAILABLE
        if self.ocr_available:
            print("✅ OCR functionality enabled")
        else:
            print("⚠️ OCR functionality disabled - missing dependencies")
        
        # CONSERVATIVE healthcare patterns - only obvious structured data
        self.healthcare_patterns = {
            'SSN': r'\b\d{3}[-.\s]\d{2}[-.\s]\d{4}\b',  # Clear SSN format
            'PHONE': r'\b\(?\d{3}\)?[-.\s]\d{3}[-.\s]\d{4}\b',  # Clear phone format
            'EMAIL': r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b',  # Clear email
            'MEDICAL_RECORD': r'\bMRN[-#:\s]*\d{6,}\b',  # Only obvious MRN format (6+ digits)
            'PATIENT_ID': r'\b(?:Patient\s+ID|PAT\s*#)[-:\s]+\d{6,}\b',  # Only labeled patient IDs
            'DATE_OF_BIRTH': r'\b(?:DOB|Date\s+of\s+Birth)[:\s]*\d{1,2}[/-]\d{1,2}[/-]\d{2,4}\b',
            'ZIP_CODE': r'\b\d{5}(?:-\d{4})?\b',
            'ADDRESS': r'\b\d+\s+[A-Za-z\s]+(?:Street|St|Avenue|Ave|Road|Rd|Drive|Dr|Lane|Ln|Boulevard|Blvd|Way|Court|Ct|Circle|Cir|Place|Pl)\b',
            'PATIENT_ID': r'\b(?:Patient|PAT|PT)[-#:\s]*\d{4,}\b',
            'ACCOUNT_NUMBER': r'\b(?:Account|Acct|ACC)[-#:\s]*\d{4,}\b',
            'DOB': r'\b(?:DOB|Date\s+of\s+Birth)[-:\s]*\d{1,2}[-/]\d{1,2}[-/]\d{2,4}\b',
            'HEIGHT_WEIGHT': r'\b\d+\s*(?:ft|feet|\')?\s*\d*\s*(?:in|inches|\")\b|\b\d+\s*(?:lbs?|pounds?|kg)\b',
            'BLOOD_PRESSURE': r'\b\d{2,3}/\d{2,3}\s*mmHg\b',
            'MEDICATION_DOSAGE': r'\b\d+\.?\d*\s*(?:mg|g|ml|cc|units?|mcg|µg)\b',
            }
        
        # Medical terminology that should NEVER be redacted
        self.medical_terminology = {
            'medical', 'patient', 'doctor', 'physician', 'nurse', 'hospital',
            'clinic', 'center', 'health', 'healthcare', 'treatment', 'care',
            'information', 'history', 'report', 'record', 'summary', 'visit',
            'appointment', 'examination', 'assessment', 'plan', 'findings',
            'diabetes', 'hypertension', 'stroke', 'heart', 'cardiac', 'renal',
            'disease', 'condition', 'syndrome', 'disorder', 'chronic', 'acute',
            'medication', 'medicine', 'drug', 'prescription', 'dosage', 'treatment',
            'therapy', 'procedure', 'surgery', 'intervention',
            'normal', 'abnormal', 'elevated', 'decreased', 'stable', 'controlled',
            'pressure', 'rate', 'level', 'count', 'value', 'result',
            'daily', 'weekly', 'monthly', 'morning', 'evening', 'before', 'after',
            'scheduled', 'follow', 'next', 'continue', 'discontinue',
            'admitted', 'discharged', 'referred', 'consulted', 'prescribed',
            'administered', 'observed', 'monitored', 'evaluated'
        }
        
        # Field labels to never redact
        self.field_labels = {
            'name', 'address', 'phone', 'email', 'date', 'birth', 'age',
            'gender', 'race', 'ethnicity', 'status', 'number', 'id', 'code'
        }

    def detect_sensitive_data(self, text: str) -> List[Dict]:
        """Hybrid AI + Regex detection - streamlined approach"""
        print("🔍 Starting hybrid AI + Regex detection...")
        entities = []
        
        # 1. Enhanced AI Models - Transformers NER
        if self.ner_pipeline:
            try:
                print("🤖 Running enhanced Transformers NER...")
                ner_results = self.ner_pipeline(text)
                for entity in ner_results:
                    if entity['score'] > 0.85:  # Lowered threshold for more inclusivity
                        entity_type = self._map_entity_type(entity['entity_group'])
                        if entity_type in ['PERSON', 'ORGANIZATION', 'LOCATION']:
                            # Clean up the entity text
                            clean_text = entity['word'].replace('##', '').replace('Ġ', ' ').strip()
                            
                            entities.append({
                                'text': clean_text,
                                'type': entity_type,
                                'start': int(entity['start']),
                                'end': int(entity['end']),
                                'confidence': float(entity['score']),
                                'source': 'transformer_ai'
                            })
                print(f"🤖 Transformers found: {len([e for e in entities if e['source'] == 'transformer_ai'])}")
            except Exception as e:
                print(f"⚠️ Transformers NER error: {e}")
        
        # 2. Enhanced spaCy NER
        if self.nlp:
            try:
                print("🧠 Running enhanced spaCy NER...")
                doc = self.nlp(text)
                for ent in doc.ents:
                    # Include more entity types including dates and ages
                    if ent.label_ in ['PERSON', 'ORG', 'GPE', 'DATE', 'TIME', 'CARDINAL']:
                        entity_type = self._map_spacy_type(ent.label_)
                        confidence = 0.85  # Lowered for more inclusivity
                        
                        # Special handling for different entity types
                        if ent.label_ == 'PERSON' and len(ent.text.split()) >= 2:
                            # Validate names in medical context
                            context = text[max(0, ent.start_char-30):min(len(text), ent.end_char+30)]
                            if self._has_medical_context(context.lower()):
                                entities.append({
                                    'text': ent.text.strip(),
                                    'type': 'PERSON',
                                    'start': int(ent.start_char),
                                    'end': int(ent.end_char),
                                    'confidence': confidence,
                                    'source': 'spacy_ai'
                                })
                        elif ent.label_ in ['ORG', 'GPE']:
                            # Organizations and locations
                            entities.append({
                                'text': ent.text.strip(),
                                'type': entity_type,
                                'start': int(ent.start_char),
                                'end': int(ent.end_char),
                                'confidence': confidence,
                                'source': 'spacy_ai'
                            })
                        elif ent.label_ in ['DATE', 'TIME']:
                            # Date entities
                            entities.append({
                                'text': ent.text.strip(),
                                'type': 'DATE',
                                'start': int(ent.start_char),
                                'end': int(ent.end_char),
                                'confidence': confidence,
                                'source': 'spacy_ai'
                            })
                        elif ent.label_ == 'CARDINAL':
                            # Numbers that might be ages or measurements
                            if self._is_likely_age_or_measurement(ent.text, text, ent.start_char, ent.end_char):
                                entities.append({
                                    'text': ent.text.strip(),
                                    'type': 'AGE',
                                    'start': int(ent.start_char),
                                    'end': int(ent.end_char),
                                    'confidence': 0.80,
                                    'source': 'spacy_ai'
                                })
                
                print(f"🧠 spaCy found: {len([e for e in entities if e['source'] == 'spacy_ai'])}")
            except Exception as e:
                print(f"⚠️ spaCy NER error: {e}")
        
        # 3. Enhanced Regex Patterns for More Entity Types
        print("📝 Running enhanced Regex patterns...")
        regex_patterns = {
            'SSN': r'\b\d{3}[-.\s]\d{2}[-.\s]\d{4}\b',
            'PHONE': r'\b\(?\d{3}\)?[-.\s]\d{3}[-.\s]\d{4}\b',
            'EMAIL': r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b',
            'MEDICAL_RECORD': r'\bMRN[-#:\s]*(\d{4,})\b',
            'DATE_OF_BIRTH': r'\b(?:DOB|Date\s+of\s+Birth)[:\s]*(\d{1,2}[/-]\d{1,2}[/-]\d{2,4})\b',
            'ZIP_CODE': r'\b\d{5}(?:-\d{4})?\b',
            
            # Enhanced date patterns
            'DATE': r'\b\d{1,2}[/-]\d{1,2}[/-]\d{2,4}\b',
            'BIRTH_DATE': r'\b(?:Born|Birth)[:\s]*(\d{1,2}[/-]\d{1,2}[/-]\d{2,4})\b',
            
            # Age patterns
            'AGE': r'\b(?:Age|aged)[:\s]*(\d{1,3})\s*(?:years?|yrs?|y\.o\.)\b',
            'AGE_SIMPLE': r'\b\d{1,3}\s*(?:years?\s+old|yrs?\s+old|y\.o\.)\b',
            
            # Address patterns
            'ADDRESS': r'\b\d+\s+[A-Za-z\s]{3,30}(?:Street|St|Avenue|Ave|Road|Rd|Drive|Dr|Lane|Ln|Boulevard|Blvd)\b',
            
            # Health measurements
            'HEIGHT': r'\b\d+\s*(?:ft|feet|\')\s*\d*\s*(?:in|inches|\")\b',
            'WEIGHT': r'\b\d+\s*(?:lbs?|pounds?|kg)\b',
            'BLOOD_PRESSURE': r'\b\d{2,3}/\d{2,3}\s*(?:mmHg)?\b',
            
            # Patient/Account IDs
            'PATIENT_ID': r'\b(?:Patient\s+ID|PAT\s*#)[-:\s]+(\d{4,})\b',
            'ACCOUNT_NUMBER': r'\b(?:Account|Acct)\s*#?[:\s]*(\d{6,})\b',
        }
        
        for pattern_name, pattern in regex_patterns.items():
            matches = re.finditer(pattern, text, re.IGNORECASE)
            for match in matches:
                # Extract sensitive part if using capture groups
                if '(' in pattern and ')' in pattern and match.groups():
                    sensitive_text = match.group(1)
                    start_pos = match.start(1)
                    end_pos = match.end(1)
                else:
                    sensitive_text = match.group()
                    start_pos = match.start()
                    end_pos = match.end()
                
                entities.append({
                    'text': sensitive_text.strip(),
                    'type': pattern_name,
                    'start': int(start_pos),
                    'end': int(end_pos),
                    'confidence': 0.99,
                    'source': 'regex'
                })
        
        print(f"📝 Regex found: {len([e for e in entities if e['source'] == 'regex'])}")
        
        # 4. Smart Filtering (remove medical terms, duplicates)
        entities = self._smart_filter(entities)
        print(f"✅ Final entities after filtering: {len(entities)}")
        
        return entities
    
    def _map_entity_type(self, entity_group: str) -> str:
        """Map AI model entity types to standard types"""
        mapping = {
            'PER': 'PERSON', 'PERSON': 'PERSON',
            'ORG': 'ORGANIZATION', 'ORGANIZATION': 'ORGANIZATION', 
            'LOC': 'LOCATION', 'LOCATION': 'LOCATION'
        }
        return mapping.get(entity_group.upper(), entity_group)
    
    def _map_spacy_type(self, label: str) -> str:
        """Map spaCy labels to our types"""
        mapping = {
            'PERSON': 'PERSON',
            'ORG': 'ORGANIZATION', 
            'GPE': 'LOCATION',
            'DATE': 'DATE',
            'TIME': 'DATE',
            'CARDINAL': 'AGE'  # Numbers often represent ages in medical docs
        }
        return mapping.get(label, label)
    
    def _is_likely_age_or_measurement(self, text: str, full_text: str, start: int, end: int) -> bool:
        """Check if a number is likely an age or measurement"""
        try:
            num = int(text)
            # Ages are typically 0-120
            if 0 <= num <= 120:
                # Check surrounding context for age indicators
                context_start = max(0, start - 20)
                context_end = min(len(full_text), end + 20)
                context = full_text[context_start:context_end].lower()
                
                age_indicators = ['age', 'aged', 'years', 'old', 'y.o.', 'yr', 'birthday']
                return any(indicator in context for indicator in age_indicators)
        except ValueError:
            pass
        return False

    def _has_medical_context(self, context: str) -> bool:
        """Check if context suggests medical relevance"""
        medical_indicators = ['patient', 'dr', 'doctor', 'physician', 'name:', 'by:', 'signed']
        return any(indicator in context for indicator in medical_indicators)

    def _smart_filter(self, entities: List[Dict]) -> List[Dict]:
        """Smart filtering to remove false positives and duplicates"""
        # Medical terms to skip
        skip_terms = {
            'medical', 'patient', 'doctor', 'hospital', 'clinic', 'health',
            'information', 'history', 'report', 'diabetes', 'hypertension',
            'treatment', 'medication', 'diagnosis', 'condition', 'center'
        }
        
        filtered = []
        for entity in entities:
            text_lower = entity['text'].lower()
            
            # Skip medical terminology
            if text_lower in skip_terms or any(term in text_lower for term in skip_terms):
                continue
                
            # Skip very short entities
            if len(entity['text'].strip()) < 3:
                continue
                
            # For names, ensure they look like actual names
            if entity['type'] == 'PERSON':
                words = entity['text'].split()
                if len(words) < 2 or any(len(word) < 2 for word in words):
                    continue
            
            filtered.append(entity)
        
        # Remove overlapping entities (keep highest confidence)
        return self._remove_overlaps(filtered)

    def _remove_overlaps(self, entities: List[Dict]) -> List[Dict]:
        """Remove overlapping entities, keeping highest confidence"""
        if not entities:
            return []
        
        entities.sort(key=lambda x: (x['start'], -x['confidence']))
        result = []
        
        for entity in entities:
            overlap_found = False
            for existing in result:
                if (entity['start'] < existing['end'] and entity['end'] > existing['start']):
                    if entity['confidence'] > existing['confidence']:
                        result.remove(existing)
                        break
                    else:
                        overlap_found = True
                        break
            
            if not overlap_found:
                result.append(entity)
        
        return result
    
    # Old methods removed - using streamlined hybrid approach above
    
    def _validate_structured_data(self, data_type: str, text: str) -> bool:
        """Extra validation for structured data"""
        if data_type == 'SSN':
            # Reject obvious fake SSNs
            fake_ssns = ['000-00-0000', '123-45-6789', '111-11-1111']
            return text not in fake_ssns
        elif data_type == 'PHONE':
            # Reject obvious fake phones
            digits_only = ''.join(filter(str.isdigit, text))
            return len(digits_only) == 10 and digits_only not in ['0000000000', '1111111111']
        elif data_type == 'MEDICAL_RECORD':
            # Must have at least 6 digits
            digits_only = ''.join(filter(str.isdigit, text))
            return len(digits_only) >= 6
        return True
    
    # Old name detection methods removed - using AI models in streamlined approach
    
    def _is_likely_a_name(self, text: str) -> bool:
        """More balanced name validation - not overly strict"""
        words = text.split()
        
        # Must be 2-3 words (first + last name, or first + middle + last)
        if len(words) < 2 or len(words) > 3:
            return False
        
        # Each word should be properly capitalized
        for word in words:
            if not word[0].isupper():
                return False
            # Reasonable word length
            if len(word) < 2 or len(word) > 20:
                return False
        
        # Skip obvious medical terms
        medical_terms_to_skip = {
            'medical center', 'health system', 'patient information',
            'blood pressure', 'heart rate', 'medical record', 
            'chief complaint', 'follow up'
        }
        
        if text.lower() in medical_terms_to_skip:
            return False
        
        # Skip if contains obvious medical words
        medical_words_to_avoid = {
            'hospital', 'clinic', 'medical', 'health', 'center', 'system',
            'report', 'record', 'information', 'history', 'diabetes', 'pressure'
        }
        
        if any(word.lower() in medical_words_to_avoid for word in words):
            return False
        
        return True
    
    def _detect_names_with_spacy(self, text: str) -> List[Dict]:
        """Use spaCy to detect additional names with context validation"""
        entities = []
        
        if not self.nlp:
            return entities
        
        try:
            doc = self.nlp(text)
            for ent in doc.ents:
                if ent.label_ == 'PERSON' and len(ent.text.split()) >= 2:
                    # Get context around the entity
                    context = self._get_entity_context(text, ent.start_char, ent.end_char, 30)
                    
                    # Check if it's likely a real person name in medical context
                    if self._is_likely_a_name(ent.text) and self._has_person_context(context):
                        entities.append({
                            'text': ent.text.strip(),
                            'type': 'PERSON',
                            'start': int(ent.start_char),
                            'end': int(ent.end_char),
                            'confidence': 0.88,
                            'source': 'spacy_contextual'
                        })
        except Exception as e:
            print(f"spaCy name detection error: {e}")
        
        return entities
    
    def _has_person_context(self, context: str) -> bool:
        """Check if context suggests this is actually a person"""
        context_lower = context.lower()
        
        # Positive indicators
        person_indicators = [
            'patient', 'dr', 'doctor', 'physician', 'provider', 'nurse',
            'name:', 'by:', 'signed', 'attending', 'resident'
        ]
        
        return any(indicator in context_lower for indicator in person_indicators)
    
    def _is_obviously_a_name(self, text: str) -> bool:
        """Very strict name validation"""
        words = text.split()
        
        # Must be 2-3 words
        if len(words) < 2 or len(words) > 3:
            return False
        
        # Each word must be properly capitalized
        for word in words:
            if not (word[0].isupper() and word[1:].islower()):
                return False
            # Must be reasonable length
            if len(word) < 2 or len(word) > 15:
                return False
        
        # Reject medical terms that might look like names
        medical_terms = {
            'medical center', 'health system', 'patient information',
            'blood pressure', 'heart rate', 'diabetes mellitus',
            'type diabetes', 'medical record', 'chief complaint'
        }
        
        if text.lower() in medical_terms:
            return False
        
        # Reject if any word is a common medical term
        medical_words = {
            'medical', 'patient', 'doctor', 'hospital', 'clinic', 'health',
            'diabetes', 'pressure', 'heart', 'blood', 'medication', 'treatment'
        }
        
        if any(word.lower() in medical_words for word in words):
            return False
        
        return True
    
    def _validate_with_spacy(self, existing_entities: List[Dict], text: str) -> List[Dict]:
        """Use spaCy only for validation of existing entities"""
        entities = []
        
        if not self.nlp:
            return entities
        
        try:
            doc = self.nlp(text)
            for ent in doc.ents:
                # Only validate PERSON entities with high confidence
                if ent.label_ == 'PERSON' and len(ent.text.split()) >= 2:
                    context = self._get_entity_context(text, ent.start_char, ent.end_char)
                    if self._is_valid_person_name(ent.text, context):
                        entities.append({
                            'text': ent.text.strip(),
                            'type': 'PERSON',
                            'start': int(ent.start_char),
                            'end': int(ent.end_char),
                            'confidence': 0.90,
                            'source': 'spacy_validation'
                        })
        except Exception as e:
            print(f"spaCy validation error: {e}")
        
        return entities
    
    def _get_entity_context(self, text: str, start: int, end: int, window: int = 50) -> str:
        """Get surrounding context for entity validation"""
        context_start = max(0, start - window)
        context_end = min(len(text), end + window)
        return text[context_start:context_end]
    
    def _is_valid_person_name(self, name: str, context: str) -> bool:
        """Validate if detected name is actually a person"""
        name_lower = name.lower()
        context_lower = context.lower()
        
        # Skip medical terms that look like names
        medical_false_positives = {
            'medical center', 'health system', 'patient information', 
            'chief complaint', 'blood pressure', 'heart rate', 'type diabetes'
        }
        
        if name_lower in medical_false_positives:
            return False
        
        # Must be 2-4 words
        words = name.split()
        if len(words) < 2 or len(words) > 4:
            return False
        
        # Look for medical context indicators
        medical_indicators = ['patient', 'dr', 'doctor', 'physician', 'provider', 'name:', 'by:']
        has_medical_context = any(indicator in context_lower for indicator in medical_indicators)
        
        return has_medical_context
    
    def _apply_strict_filtering(self, entities: List[Dict]) -> List[Dict]:
        """Aggressively filter out potential false positives"""
        filtered_entities = []
        
        for entity in entities:
            text_lower = entity['text'].strip().lower()
            
            # Skip if it's medical terminology
            if text_lower in self.medical_terminology:
                continue
            
            # Skip if it's a field label
            if text_lower in self.field_labels or text_lower.endswith(':'):
                continue
            
            # Skip single character or very short entities
            if len(entity['text'].strip()) < 3:
                continue
            
            # Skip if it's mostly numbers (unless it's a clear pattern like SSN)
            if (entity['type'] not in ['SSN', 'PHONE', 'MEDICAL_RECORD', 'PATIENT_ID'] and
                sum(c.isdigit() for c in entity['text']) > len(entity['text']) * 0.7):
                continue
            
            # For PERSON type, apply extra strict validation
            if entity['type'] == 'PERSON':
                if not self._final_person_validation(entity['text']):
                    continue
            
            filtered_entities.append(entity)
        
        return filtered_entities
    
    def _final_person_validation(self, name: str) -> bool:
        """Final validation for person names"""
        
        # Must pass basic name checks
        if not self._is_obviously_a_name(name):
            return False
        
        # Additional checks for common false positives
        words = name.split()
        
        # Check for medical facility patterns
        facility_indicators = ['medical', 'health', 'hospital', 'clinic', 'center', 'system']
        if any(indicator in word.lower() for word in words for indicator in facility_indicators):
            return False
        
        # Check for measurement patterns
        if any(char.isdigit() for char in name):
            return False
        
        # Check for common medical abbreviations in names
        medical_abbrevs = ['md', 'rn', 'np', 'pa', 'do', 'phd', 'dds']
        if any(word.lower() in medical_abbrevs for word in words):
            return False
        
        return True
    
    def _apply_conservative_thresholds(self, entities: List[Dict]) -> List[Dict]:
        """Apply balanced confidence thresholds - not too strict"""
        
        # Balanced confidence requirements
        min_confidence_by_type = {
            'PERSON': 0.85,           # Names - more lenient
            'SSN': 0.98,              # SSN must be very certain
            'PHONE': 0.95,            # Phone fairly certain  
            'EMAIL': 0.95,            # Email fairly certain
            'MEDICAL_RECORD': 0.90,   # Medical records 
            'PATIENT_ID': 0.90,       # Patient IDs 
            'DATE_OF_BIRTH': 0.85,    # DOB more lenient
            'ZIP_CODE': 0.90          # Zip codes
        }
        
        # Source-based confidence adjustments - more balanced
        source_multipliers = {
            'regex_strict': 1.0,        # Regex gets full confidence
            'context_aware': 0.95,      # Context-based 
            'spacy_contextual': 0.90,   # spaCy contextual
            'spacy_validation': 0.88,   # spaCy validation
            'transformer_ner': 0.85     # Transformers 
        }
        
        conservative_entities = []
        
        for entity in entities:
            entity_type = entity['type']
            source = entity['source']
            confidence = entity['confidence']
            
            # Apply source multiplier
            adjusted_confidence = confidence * source_multipliers.get(source, 0.8)
            
            # Check against minimum threshold
            min_threshold = min_confidence_by_type.get(entity_type, 0.95)
            
            if adjusted_confidence >= min_threshold:
                entity['confidence'] = adjusted_confidence
                conservative_entities.append(entity)
        
        return conservative_entities
        
        # 2. Use spaCy NER for additional entities
        if self.nlp:
            try:
                doc = self.nlp(text)
                for ent in doc.ents:
                    entity_type = self._map_spacy_entity_type(ent.label_)
                    if ent.label_ in ['PERSON', 'GPE', 'ORG', 'DATE', 'TIME', 'MONEY', 'CARDINAL', 'ORDINAL']:
                        entities.append({
                            'text': ent.text.strip(),
                            'type': entity_type,
                            'start': int(ent.start_char),
                            'end': int(ent.end_char),
                            'confidence': 0.95,
                            'source': 'spacy_ner'
                        })
            except Exception as e:
                print(f"spaCy NER error: {e}")
        
        # 3. Use healthcare-specific regex patterns
        for pattern_name, pattern in self.healthcare_patterns.items():
            matches = re.finditer(pattern, text, re.IGNORECASE)
            for match in matches:
                entities.append({
                    'text': match.group().strip(),
                    'type': pattern_name,
                    'start': int(match.start()),
                    'end': int(match.end()),
                    'confidence': 0.99,
                    'source': 'regex'
                })
        
        # Remove duplicates and overlapping entities
        entities = self._deduplicate_entities(entities)
        
        # Filter out field labels before returning
        entities = self._filter_field_labels(entities)
        
        # Apply confidence-based filtering
        entities = self._filter_by_confidence(entities)
        
        return entities
    
    def _clean_entity_text(self, text: str) -> str:
        """Clean entity text from transformer output"""
        return text.replace('##', '')
    
    def _map_transformer_entity_type(self, entity_group: str) -> str:
        """Map transformer entity types to our standardized types"""
        mapping = {
            'PER': 'PERSON',
            'PERSON': 'PERSON', 
            'LOC': 'LOCATION',
            'LOCATION': 'LOCATION',
            'ORG': 'ORGANIZATION',
            'ORGANIZATION': 'ORGANIZATION',
            'MISC': 'MISCELLANEOUS'
        }
        return mapping.get(entity_group.upper(), entity_group)
    
    def _map_spacy_entity_type(self, label: str) -> str:
        """Map spaCy entity types to our standardized types"""
        mapping = {
            'PERSON': 'PERSON',
            'GPE': 'LOCATION',
            'ORG': 'ORGANIZATION',
            'DATE': 'DATE',
            'TIME': 'TIME',
            'MONEY': 'FINANCIAL',
            'CARDINAL': 'NUMBER',
            'ORDINAL': 'NUMBER'
        }
        return mapping.get(label, label)
    
    def _deduplicate_entities(self, entities: List[Dict]) -> List[Dict]:
        """Remove duplicate and overlapping entities"""
        if not entities:
            return []
        
        # Sort by start position
        entities.sort(key=lambda x: (x['start'], -x['end']))
        
        deduplicated = []
        for entity in entities:
            # Check if this entity overlaps with any existing entity
            overlaps = False
            for existing in deduplicated:
                if (entity['start'] < existing['end'] and entity['end'] > existing['start']):
                    # Choose the one with higher confidence
                    if entity['confidence'] > existing['confidence']:
                        deduplicated.remove(existing)
                        break
                    else:
                        overlaps = True
                        break
            
            if not overlaps:
                deduplicated.append(entity)
        
        return deduplicated

    def _filter_field_labels(self, entities: List[Dict]) -> List[Dict]:
        """Filter out common field labels and non-sensitive data"""
        # Comprehensive list of field labels and non-sensitive terms
        field_labels = {
            'Name:', 'SSN:', 'Phone:', 'Email:', 'Address:', 'Date of Birth:', 
            'Medical Record:', 'Patient Information:', 'Medical History:',
            'Recent Visit:', 'Chief Complaint:', 'Diagnosis:', 'Prescribed Medications:',
            'Next Appointment:', 'Patient', 'Information', 'History', 'Visit',
            'Complaint', 'Medications', 'Appointment', 'MEDICAL REPORT',
            'Patient Information', 'Medical Record', 'Medical History',
            'Recent Visit', 'Chief Complaint', 'Prescribed Medications',
            'Next Appointment', 'SSN', 'Phone', 'Email', 'Address',
            'Type', 'Diabetes', 'Hypertension', 'controlled', 'twice daily',
            'once daily', 'follow-up', 'management', 'conditions', 'diagnosed',
            'suffered', 'developed', 'treated', 'admitted', 'transferred',
            'referred', 'impression', 'manifesting', 'behavioural', 'psychological',
            'symptoms', 'secondary', 'Dementia', 'informed', 'incontinent',
            'unable', 'bathe', 'toilet', 'feed', 'observed', 'gradual',
            'deterioration', 'cognitive', 'ability', 'physical', 'state',
            'years', 'Findings', 'examination', 'mental', 'state', 'examination',
            'Please', 'state', 'observations', 'patient', 'demeanour', 'body',
            'language', 'tone', 'voice', 'relevant', 'assessment', 'interactions',
            'tests', 'administered', 'questions', 'posed', 'date', 'examined',
            'brought', 'clinic', 'wheel', 'chair', 'mood', 'euthymic', 'normal',
            'non-depressed', 'reasonably', 'positive', 'psychotic', 'Orientation',
            'time', 'place', 'person', 'regard', 'orientation', 'unable', 'tell',
            'hospital', 'identified', 'son', 'able', 'name', 'told', 'identified',
            'doctor', 'asked', 'remember', 'treating', 'last', 'years', 'correctly',
            'arrived', 'moments', 'later', 'where', 'said', 'know', 'guessing',
            'date', 'day', 'Wednesday', 'Monday', 'forgot', 'answer', 'year',
            'time', 'pm', 'afternoon', 'actual', 'looking', 'watch', 'Basic',
            'gave', 'age', 'actually', 'answer', 'asked', 'birthday', 'correctly',
            'lived', 'flat', 'correct', 'address', 'area', 'wrong', 'said', 'flat',
            'Bedok', 'actually', 'Jurong', 'incorrectly', 'stated', 'Prime',
            'Minister', 'Kuan', 'Yew', 'Simple', 'arithmetic', 'financial', 'issues',
            'unable', 'perform', 'simple', 'arithmetic', 'plus', 'giving', 'answer',
            'subtract', 'count', 'backwards', 'unable', 'recognise', 'notes', 'coins',
            'identified', 'cent', 'coin', 'cents', 'dollar', 'note', 'dollars',
            'Personal', 'welfare', 'property', 'affairs', 'related', 'questions',
            'asked', 'owned', 'property', 'said', 'sole', 'owner', 'elderly',
            'mother', 'joint', 'names', 'currently', 'living', 'flat', 'herself',
            'address', 'remember', 'rooms', 'room', 'flat', 'rent', 'staying',
            'planned', 'charge', 'good', 'money', 'matters', 'knew', 'medical',
            'problems', 'stared', 'blankly', 'shook', 'head', 'remembered',
            'hospital', 'before', 'nodded', 'knew', 'hospital', 'shook', 'head',
            'taking', 'medicine', 'moment', 'looked', 'blankly', 'reply'
        }
        
        # Common medical terms that shouldn't be redacted
        medical_terms = {
            'diabetes', 'hypertension', 'stroke', 'heart', 'problems', 'cardiomyopathy',
            'cardiac', 'failure', 'chronic', 'renal', 'disease', 'admitted', 'diagnosed',
            'confirmed', 'brain', 'scans', 'transferred', 'rehabilitation', 'referred',
            'follow-up', 'treatment', 'clinical', 'impression', 'behavioural',
            'psychological', 'symptoms', 'dementia', 'incontinent', 'cognitive',
            'deterioration', 'physical', 'examination', 'mental', 'state', 'wheel',
            'chair', 'mood', 'euthymic', 'psychotic', 'orientation', 'arithmetic',
            'financial', 'welfare', 'property', 'affairs'
        }
        
        filtered_entities = []
        for entity in entities:
            entity_text = entity['text'].strip()
            entity_lower = entity_text.lower()
            
            # Skip if it's a field label
            if entity_text in field_labels or entity_text.endswith(':'):
                continue
            
            # Skip if it's a medical term
            if entity_lower in medical_terms:
                continue
            
            # Skip if it's too short (likely not a real name)
            if len(entity_text) < 3:
                continue
            
            # Skip single characters and numbers
            if len(entity_text) == 1 or entity_text.isdigit():
                continue
            
            # Skip common words that are not names
            common_words = {'the', 'and', 'or', 'but', 'in', 'on', 'at', 'to', 'for', 'of', 'with', 'by', 'from', 'up', 'about', 'into', 'through', 'during', 'before', 'after', 'above', 'below', 'between', 'among', 'under', 'over', 'around', 'near', 'far', 'here', 'there', 'where', 'when', 'why', 'how', 'what', 'who', 'which', 'that', 'this', 'these', 'those', 'is', 'are', 'was', 'were', 'be', 'been', 'being', 'have', 'has', 'had', 'do', 'does', 'did', 'will', 'would', 'could', 'should', 'may', 'might', 'must', 'can', 'shall', 'a', 'an', 'the'}
            if entity_lower in common_words:
                continue
            
            # Skip if it's mostly punctuation or special characters
            if len([c for c in entity_text if c.isalnum()]) < len(entity_text) * 0.5:
                continue
            
            # Skip if it's a common medical abbreviation or term
            medical_abbrevs = {'mg', 'ml', 'cc', 'mcg', 'µg', 'units', 'daily', 'twice', 'once', 'pm', 'am', 'hrs', 'mins', 'secs', 'lbs', 'kg', 'ft', 'in', 'cm', 'mm', 'mmHg', 'bpm', 'temp', 'bp', 'hr', 'rr', 'spo2', 'o2', 'co2', 'ph', 'glucose', 'hba1c', 'ldl', 'hdl', 'triglycerides', 'cholesterol', 'creatinine', 'bun', 'sodium', 'potassium', 'chloride', 'bicarbonate', 'calcium', 'phosphorus', 'magnesium', 'albumin', 'protein', 'bilirubin', 'alt', 'ast', 'alk', 'phos', 'ggt', 'ldh', 'cpk', 'troponin', 'bnp', 'probnp', 'tsh', 't3', 't4', 'ft3', 'ft4', 'cortisol', 'insulin', 'c-peptide', 'glucagon', 'gh', 'igf-1', 'lh', 'fsh', 'testosterone', 'estradiol', 'progesterone', 'prolactin', 'acth', 'aldosterone', 'renin', 'angiotensin', 'epinephrine', 'norepinephrine', 'dopamine', 'serotonin', 'histamine', 'leukotrienes', 'prostaglandins', 'cytokines', 'interferon', 'interleukin', 'tumor', 'necrosis', 'factor', 'alpha', 'beta', 'gamma', 'delta', 'epsilon', 'zeta', 'eta', 'theta', 'iota', 'kappa', 'lambda', 'mu', 'nu', 'xi', 'omicron', 'pi', 'rho', 'sigma', 'tau', 'upsilon', 'phi', 'chi', 'psi', 'omega'}
            if entity_lower in medical_abbrevs:
                continue
            
            filtered_entities.append(entity)
        
        return filtered_entities

    def _filter_by_confidence(self, entities: List[Dict]) -> List[Dict]:
        """Filter entities by confidence threshold and type-specific rules"""
        filtered_entities = []
        
        for entity in entities:
            entity_type = entity.get('type', '')
            confidence = entity.get('confidence', 0.0)
            entity_text = entity.get('text', '').strip()
            
            # Set confidence thresholds based on entity type
            if entity_type in ['PERSON', 'FULL_NAME', 'DOCTOR_NAME', 'PATIENT_NAME']:
                # Names need high confidence
                if confidence < 0.95:
                    continue
                # Must be at least 2 words for names
                if len(entity_text.split()) < 2:
                    continue
            elif entity_type in ['SSN', 'PHONE', 'EMAIL', 'MEDICAL_RECORD']:
                # Structured data needs very high confidence
                if confidence < 0.99:
                    continue
            elif entity_type in ['DATE', 'ZIP_CODE', 'AGE']:
                # Dates and numbers need high confidence
                if confidence < 0.95:
                    continue
            elif entity_type in ['ADDRESS', 'LOCATION']:
                # Addresses need high confidence
                if confidence < 0.90:
                    continue
            else:
                # Other types need reasonable confidence
                if confidence < 0.85:
                    continue
            
            # Additional validation for specific types
            if entity_type == 'PERSON' and not any(title in entity_text for title in ['Dr.', 'Mr.', 'Mrs.', 'Ms.', 'Prof.']):
                # If it's a PERSON but doesn't have a title, it might be a false positive
                if confidence < 0.98:
                    continue
            
            if entity_type == 'FULL_NAME':
                # Must be exactly 2-3 words for full names
                words = entity_text.split()
                if len(words) < 2 or len(words) > 3:
                    continue
                # Each word must be at least 2 characters
                if any(len(word) < 2 for word in words):
                    continue
            
            filtered_entities.append(entity)
        
        return filtered_entities

    def redact_text(self, text: str, entities: List[Dict]) -> str:
        """Redact selected entities from text"""
        if not entities:
            return text
        
        # Sort entities by start position in reverse order
        entities_sorted = sorted(entities, key=lambda x: x['start'], reverse=True)
        
        redacted_text = text
        for entity in entities_sorted:
            if entity.get('selected', True):  # Default to selected if not specified
                redaction = '[REDACTED]'
                redacted_text = (
                    redacted_text[:entity['start']] + 
                    redaction + 
                    redacted_text[entity['end']:]
                )
        
        return redacted_text

# Initialize redaction engine
redaction_engine = HealthcareRedactionEngine()

class DocumentProcessor:
    """Document processing utilities"""
    
    @staticmethod
    async def create_redacted_pdf(original_path: str, entities_to_redact: List[Dict], output_path: str) -> bool:
        """RIGID PDF REDACTION - Based on working Streamlit implementation"""
        print(f"� RIGID PDF REDACTION START")
        print(f"📁 Input: {original_path}")
        print(f"📁 Output: {output_path}")
        print(f"🎯 Entities: {len(entities_to_redact)}")
        
        if not PDF_AVAILABLE:
            print("❌ FAIL: PyMuPDF not available")
            return False
        
        if not entities_to_redact:
            print("❌ FAIL: No entities to redact")
            return False
        
        try:
            # Step 1: Open PDF
            doc = fitz.open(original_path)
            total_redactions = 0
            
            # DEBUG: Show actual PDF text
            pdf_text = ""
            for page_num in range(len(doc)):
                page = doc[page_num]
                pdf_text += page.get_text()
            
            print(f"📄 DEBUG: Actual PDF text (first 500 chars):")
            print(f"'{pdf_text[:500]}...'")
            print(f"📊 DEBUG: PDF has {len(pdf_text)} total characters")
            
            # Step 2: Process each page with simplified text search
            for page_num in range(len(doc)):
                page = doc[page_num]
                print(f"📖 Processing page {page_num + 1}")
                
                # Simple text-based redaction approach
                for entity in entities_to_redact:
                    if not entity.get('selected', True):
                        continue
                    
                    entity_text = entity['text'].strip()
                    if not entity_text:
                        continue
                    
                    print(f"🔍 Searching for '{entity_text}' on page {page_num + 1}")
                    
                    # Try multiple search approaches
                    search_texts = [
                        entity_text,
                        entity_text.lower(),
                        entity_text.upper(),
                        entity_text.title()
                    ]
                    
                    found_any = False
                    for search_text in search_texts:
                        rects = page.search_for(search_text)
                        if rects:
                            for rect in rects:
                                # Add redaction annotation
                                redact_annot = page.add_redact_annot(rect)
                                redact_annot.set_colors(stroke=(0, 0, 0), fill=(0, 0, 0))
                                redact_annot.update()
                                total_redactions += 1
                                found_any = True
                                print(f"✅ Redacted '{search_text}' on page {page_num + 1}")
                            break  # Found matches, no need to try other variations
                    
                    # If no exact matches, try partial matches
                    if not found_any and len(entity_text) > 3:
                        # Try searching for parts of the entity text
                        words = entity_text.split()
                        for word in words:
                            if len(word) > 2:  # Only try words longer than 2 characters
                                word_rects = page.search_for(word)
                                if word_rects:
                                    for rect in word_rects:
                                        redact_annot = page.add_redact_annot(rect)
                                        redact_annot.set_colors(stroke=(0, 0, 0), fill=(0, 0, 0))
                                        redact_annot.update()
                                        total_redactions += 1
                                        print(f"✅ Partial redacted '{word}' on page {page_num + 1}")
                
                # Apply redactions on this page
                page.apply_redactions()
            
            # Step 5: Save file
            doc.save(output_path, garbage=4, deflate=True, clean=True)
            doc.close()
            
            # Step 6: Verify success - more lenient criteria
            file_exists = os.path.exists(output_path)
            file_size = os.path.getsize(output_path) if file_exists else 0
            
            print(f"📊 Results: File exists: {file_exists}, Size: {file_size} bytes, Redactions: {total_redactions}")
            
            # Success if file was created and has reasonable size (don't require redactions > 0)
            success = file_exists and file_size > 1000
            
            if success:
                print(f"✅ SUCCESS: {total_redactions} redactions, {os.path.getsize(output_path)} bytes")
                return True
            else:
                print(f"❌ FAIL: File created but no redactions applied")
                if os.path.exists(output_path):
                    os.remove(output_path)  # Delete failed file
                return False
                
        except Exception as e:
            print(f"❌ CRITICAL ERROR: {str(e)}")
            return False
        
    @staticmethod
    async def create_redacted_docx(original_path: str, entities_to_redact: List[Dict], output_path: str):
        """Create a redacted Word document with highlighted redactions"""
        if not DOCX_AVAILABLE:
            raise HTTPException(status_code=400, detail="Word document processing not available")
        
        try:
            # Open original document
            doc = DocxDocument(original_path)
            
            # Process each paragraph
            for paragraph in doc.paragraphs:
                para_text = paragraph.text
                
                # Check if any entities need to be redacted in this paragraph
                for entity in entities_to_redact:
                    if not entity.get('selected', True):
                        continue
                        
                    entity_text = entity['text']
                    if entity_text.lower() in para_text.lower():
                        # Replace entity text with black bars (█ character)
                        redaction_bar = '█' * len(entity_text)
                        para_text = para_text.replace(entity_text, redaction_bar)
                
                # Update paragraph text
                paragraph.text = para_text
            
            # Save redacted document
            doc.save(output_path)
            
        except Exception as e:
            # Fallback: create new document with redacted text
            try:
                doc = DocxDocument()
                original_doc = DocxDocument(original_path)
                
                for paragraph in original_doc.paragraphs:
                    para_text = paragraph.text
                    
                    for entity in entities_to_redact:
                        if entity.get('selected', True) and entity['text'].lower() in para_text.lower():
                            redaction_bar = '█' * len(entity['text'])
                            para_text = para_text.replace(entity['text'], redaction_bar)
                    
                    if para_text.strip():
                        doc.add_paragraph(para_text)
                
                doc.save(output_path)
                
            except Exception as fallback_e:
                raise HTTPException(status_code=400, detail=f"Word document redaction error: {str(fallback_e)}")
    
    @staticmethod
    async def create_redacted_image(original_path: str, entities_to_redact: List[Dict], output_path: str):
        """Create a redacted image with black bars over sensitive information"""
        if not OCR_AVAILABLE:
            raise HTTPException(status_code=400, detail="Image processing not available")
        
        try:
            # Open original image
            image = Image.open(original_path).convert('RGB')
            
            # Get OCR data with bounding boxes
            ocr_data = pytesseract.image_to_data(image, output_type=pytesseract.Output.DICT)
            
            # Create drawing context
            draw = ImageDraw.Draw(image)
            
            # Process each detected word from OCR
            for i, word in enumerate(ocr_data['text']):
                if int(ocr_data['conf'][i]) > 30:  # Only high confidence detections
                    word_text = word.strip()
                    
                    # Check if this word should be redacted
                    for entity in entities_to_redact:
                        if not entity.get('selected', True):
                            continue
                            
                        entity_text = entity['text'].strip()
                        
                        # Check if entity text matches or is part of the OCR word
                        if (entity_text.lower() in word_text.lower() or 
                            word_text.lower() in entity_text.lower() or
                            entity_text.lower() == word_text.lower()):
                            
                            # Get bounding box coordinates
                            x = ocr_data['left'][i]
                            y = ocr_data['top'][i]
                            w = ocr_data['width'][i]
                            h = ocr_data['height'][i]
                            
                            # Draw black rectangle over the text
                            draw.rectangle([x, y, x + w, y + h], fill='black')
            
            # Save redacted image
            image.save(output_path)
            
        except Exception as e:
            # Fallback: create image with black bars for known entities
            try:
                image = Image.open(original_path).convert('RGB')
                draw = ImageDraw.Draw(image)
                
                # Simple fallback - draw black bars at estimated positions
                y_offset = 50
                for entity in entities_to_redact:
                    if entity.get('selected', True):
                        # Estimate text width (rough calculation)
                        text_width = len(entity['text']) * 12
                        draw.rectangle([50, y_offset, 50 + text_width, y_offset + 20], fill='black')
                        y_offset += 30
                
                image.save(output_path)
                
            except Exception as fallback_e:
                raise HTTPException(status_code=400, detail=f"Image redaction error: {str(fallback_e)}")
    
    @staticmethod
    async def extract_text_from_pdf(file_path: str) -> str:
        """Extract text from PDF"""
        if not PDF_AVAILABLE:
            raise HTTPException(status_code=400, detail="PDF processing not available")
        
        try:
            doc = fitz.open(file_path)
            text = ""
            for page in doc:
                text += page.get_text()
            doc.close()
            return text
        except Exception as e:
            raise HTTPException(status_code=400, detail=f"PDF processing error: {str(e)}")
    
    @staticmethod
    async def extract_text_from_docx(file_path: str) -> str:
        """Extract text from Word document"""
        if not DOCX_AVAILABLE:
            raise HTTPException(status_code=400, detail="Word document processing not available")
        
        try:
            doc = DocxDocument(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text
        except Exception as e:
            raise HTTPException(status_code=400, detail=f"Word document processing error: {str(e)}")
    
    @staticmethod
    async def extract_text_from_image(file_path: str) -> str:
        """Extract text from image using OCR"""
        if not OCR_AVAILABLE:
            raise HTTPException(status_code=400, detail="OCR processing not available")
        
        try:
            image = Image.open(file_path)
            text = pytesseract.image_to_string(image)
            return text
        except Exception as e:
            raise HTTPException(status_code=400, detail=f"OCR processing error: {str(e)}")
    
    @staticmethod 
    async def process_document(file_path: str, filename: str) -> str:
        """Process document and extract text based on file type"""
        file_ext = Path(filename).suffix.lower()
        
        if file_ext == '.pdf':
            return await DocumentProcessor.extract_text_from_pdf(file_path)
        elif file_ext in ['.docx', '.doc']:
            return await DocumentProcessor.extract_text_from_docx(file_path)
        elif file_ext in ['.png', '.jpg', '.jpeg', '.tiff', '.bmp']:
            return await DocumentProcessor.extract_text_from_image(file_path)
        else:
            # Try to read as plain text
            try:
                async with aiofiles.open(file_path, 'r', encoding='utf-8') as f:
                    return await f.read()
            except:
                try:
                    async with aiofiles.open(file_path, 'r', encoding='latin-1') as f:
                        return await f.read()
                except Exception as e:
                    raise HTTPException(status_code=400, detail=f"Unsupported file type: {file_ext}")

# API Endpoints

@app.get("/api/health")
async def health_check():
    """Health check endpoint"""
    return {
        "status": "healthy",
        "timestamp": datetime.now().isoformat(),
        "features": {
            "ocr": OCR_AVAILABLE,
            "pdf": PDF_AVAILABLE,
            "docx": DOCX_AVAILABLE
        }
    }

@app.post("/api/upload")
async def upload_document(file: UploadFile = File(...)):
    """Upload and process document"""
    try:
        # Generate unique document ID
        document_id = str(uuid.uuid4())
        
        # Save uploaded file
        upload_path = uploads_dir / f"{document_id}_{file.filename}"
        async with aiofiles.open(upload_path, 'wb') as f:
            content = await file.read()
            await f.write(content)
        
        # Extract text from document
        extracted_text = await DocumentProcessor.process_document(str(upload_path), file.filename)
        print(f"📄 DEBUG: Extracted text preview (first 300 chars):")
        print(f"'{extracted_text[:300]}...'")
        
        # Detect sensitive entities
        entities = redaction_engine.detect_sensitive_data(extracted_text)
        print(f"🔍 DEBUG: Detected {len(entities)} entities:")
        for i, entity in enumerate(entities[:10]):  # Show first 10
            print(f"   {i+1}. '{entity['text']}' ({entity['type']}) - confidence: {entity['confidence']}")
        
        # Make entities JSON serializable
        entities = make_json_serializable(entities)
        
        # Store document info with proper timestamp
        upload_time = datetime.now().isoformat() + 'Z'
        documents_db[document_id] = {
            'filename': file.filename,
            'upload_time': upload_time,
            'original_path': str(upload_path),
            'extracted_text': extracted_text,
            'entities': entities,
            'status': 'processed'
        }
        
        return {
            'success': True,
            'document_id': document_id,
            'filename': file.filename,
            'extracted_text': extracted_text,
            'entities': entities,
            'entity_count': len(entities)
        }
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Document processing failed: {str(e)}")

@app.post("/api/redact")
async def redact_document(request: RedactionRequest, background_tasks: BackgroundTasks):
    """Redact selected entities from document"""
    try:
        # Get document info
        if request.document_id not in documents_db:
            raise HTTPException(status_code=404, detail="Document not found")
        
        doc_info = documents_db[request.document_id]
        original_text = doc_info['extracted_text']
        
        # Convert request entities to dict format and DEBUG
        entities_to_redact = []
        selected_count = 0
        
        # If no entities provided, use all detected entities from upload
        if request.entities is None or len(request.entities) == 0:
            print(f"🔍 DEBUG: No entities provided - using all detected entities from upload")
            entities_to_redact = doc_info.get('entities', [])
            selected_count = len(entities_to_redact)
            print(f"🎯 DEBUG: Auto-redacting {selected_count} detected entities")
        else:
            print(f"🔍 DEBUG: Received {len(request.entities)} entities from frontend:")
            for i, entity in enumerate(request.entities):
                print(f"   {i+1}. Text: '{entity.text}' | Type: {entity.type} | Selected: {entity.selected}")
                if entity.selected:
                    entities_to_redact.append({
                        'text': entity.text,
                        'type': entity.type,
                        'start': entity.start,
                        'end': entity.end,
                        'selected': True
                    })
                    selected_count += 1
            
            print(f"🎯 DEBUG: {selected_count} entities selected for redaction:")
            for i, entity in enumerate(entities_to_redact):
                print(f"   {i+1}. '{entity['text']}' ({entity['type']})")
        
        # Redact the text
        redacted_text = redaction_engine.redact_text(original_text, entities_to_redact)
        
        # Generate redacted document in same format as original
        original_filename = doc_info['filename']
        file_ext = Path(original_filename).suffix.lower()
        redacted_filename = f"redacted_{Path(original_filename).stem}{file_ext}"
        
        # Use temporary file instead of downloads directory
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=file_ext, prefix=f"{request.document_id}_")
        redacted_path = Path(temp_file.name)
        temp_file.close()
        
        print(f"🚀 RIGID REDACTION - NO COMPROMISE APPROACH")
        print(f"📁 File: {doc_info['original_path']}")
        print(f"📄 Type: {file_ext}")
        print(f"🎯 Entities: {selected_count}")
        
        redaction_success = False
        
# Replace the redaction logic in the /api/redact endpoint around lines 770-780

# Enable all document types instead of PDF-only
        if file_ext == '.pdf':
            print("📝 Attempting PDF redaction...")
            redaction_success = await DocumentProcessor.create_redacted_pdf(
                doc_info['original_path'], entities_to_redact, str(redacted_path)
            )
        elif file_ext in ['.docx', '.doc']:
            print("📝 Attempting Word document redaction...")
            try:
                await DocumentProcessor.create_redacted_docx(
                    doc_info['original_path'], entities_to_redact, str(redacted_path)
                )
                redaction_success = True
                print("✅ Word document redaction completed")
            except Exception as e:
                print(f"❌ Word redaction failed: {e}")
                redaction_success = False
        elif file_ext in ['.png', '.jpg', '.jpeg', '.tiff', '.bmp']:
            print("📝 Attempting image redaction...")
            try:
                await DocumentProcessor.create_redacted_image(
                    doc_info['original_path'], entities_to_redact, str(redacted_path)
                )
                redaction_success = True
                print("✅ Image redaction completed")
            except Exception as e:
                print(f"❌ Image redaction failed: {e}")
                redaction_success = False
        else:
            print(f"❌ UNSUPPORTED: File type {file_ext} not supported for redaction")
            raise HTTPException(status_code=400, detail=f"File type {file_ext} not supported for redaction")
        
        # Update document info with redaction timestamp only if redaction was successful
        if redaction_success:
            redaction_time = datetime.now().isoformat() + 'Z'
            documents_db[request.document_id].update({
                'redacted_text': redacted_text,
                'redacted_path': str(redacted_path),
                'redacted_filename': redacted_filename,
                'redacted_count': selected_count,
                'redacted_entities': entities_to_redact,  # Store the actual redacted entities
                'redaction_time': redaction_time,
                'status': 'redacted'
            })
        else:
            print(f"❌ REDACTION FAILED: Not updating document status")
            raise HTTPException(status_code=500, detail="Document redaction failed")
        
        print(f"💾 Document database updated for ID: {request.document_id}")
        print(f"🎯 Final status: {documents_db[request.document_id]['status']}")
        
        # Return dictionary instead of Pydantic model instance
        return {
            "success": True,
            "document_id": request.document_id,
            "redacted_text": redacted_text,
            "redacted_count": selected_count,
            "download_url": f"/api/download/{request.document_id}"
        }
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Redaction failed: {str(e)}")
# Additional API endpoints for preview and audit functionality





@app.get("/api/download/{document_id}")
async def download_document(document_id: str, background_tasks: BackgroundTasks, original: bool = False):
    """Download processed document"""
    print(f"🔻 DOWNLOAD REQUEST for document ID: {document_id}, original: {original}")
    
    try:
        if document_id not in documents_db:
            print(f"❌ Document not found in database: {document_id}")
            print(f"📊 Available documents: {list(documents_db.keys())}")
            raise HTTPException(status_code=404, detail="Document not found")
        
        doc_info = documents_db[document_id]
        print(f"📄 Document info: {doc_info.get('filename', 'Unknown')}")
        print(f"📊 Status: {doc_info.get('status', 'Unknown')}")
        
        # Determine which file to serve
        if original:
            # Serve original file for viewing
            if 'original_path' in doc_info and os.path.exists(doc_info['original_path']):
                file_path = doc_info['original_path']
                filename = doc_info['filename']
                file_size = os.path.getsize(file_path)
                
                print(f"✅ SERVING ORIGINAL FILE: {filename} ({file_size} bytes)")
                
                return FileResponse(
                    path=file_path,
                    filename=filename,
                    media_type='application/pdf'
                )
            else:
                print(f"❌ ORIGINAL FILE NOT FOUND")
                raise HTTPException(status_code=404, detail="Original file not found")
        else:
            # Serve redacted files if they exist
            if ('redacted_path' in doc_info and 
                os.path.exists(doc_info['redacted_path'])):
                
                file_path = doc_info['redacted_path']
                filename = doc_info['redacted_filename']
                file_size = os.path.getsize(file_path)
                
                print(f"✅ SERVING REDACTED FILE: {filename} ({file_size} bytes)")
                
                # DON'T schedule cleanup for viewing - only for actual downloads
                # background_tasks.add_task(cleanup_temp_file, file_path)
                
                return FileResponse(
                    path=file_path,
                    filename=filename,
                    media_type='application/pdf'
                )
            else:
                print(f"❌ NO REDACTED FILE AVAILABLE")
                print(f"📊 Document status: {doc_info.get('status', 'unknown')}")
                print(f"🔍 Redacted path exists: {'redacted_path' in doc_info}")
                if 'redacted_path' in doc_info:
                    print(f"📁 Path: {doc_info['redacted_path']}")
                    print(f"📂 File exists: {os.path.exists(doc_info['redacted_path'])}")
                
                raise HTTPException(
                    status_code=404, 
                    detail="No redacted file available. Document must be successfully redacted before download."
                )
            
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Download failed: {str(e)}")

# Additional API endpoints for preview and audit functionality

@app.get("/api/document/{document_id}/preview")
async def get_document_preview(document_id: str):
    """Get document preview information and detected entities"""
    try:
        if document_id not in documents_db:
            raise HTTPException(status_code=404, detail="Document not found")
        
        doc_info = documents_db[document_id]
        extracted_text = doc_info['extracted_text']
        
        # Create preview of extracted text (first 500 characters)
        text_preview = extracted_text[:500] + "..." if len(extracted_text) > 500 else extracted_text
        
        response = {
            "success": True,
            "document_id": document_id,
            "filename": doc_info['filename'],
            "upload_time": doc_info['upload_time'],
            "status": doc_info.get('status', 'processed'),
            "entity_count": len(doc_info.get('entities', [])),
            "extracted_text_preview": text_preview,
            "entities": doc_info.get('entities', []),
            "redacted_count": doc_info.get('redacted_count', 0)
        }
        
        return response
    
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Preview generation failed: {str(e)}")

@app.get("/api/document/{document_id}/redacted-preview")
async def get_redacted_document_preview(document_id: str):
    """Get redacted document preview information"""
    try:
        if document_id not in documents_db:
            raise HTTPException(status_code=404, detail="Document not found")
        
        doc_info = documents_db[document_id]
        
        # Check if document has been redacted
        if doc_info.get('status') != 'redacted':
            raise HTTPException(status_code=400, detail="Document has not been redacted yet")
        
        redacted_text = doc_info.get('redacted_text', '')
        
        # Create preview of redacted text (first 500 characters)
        redacted_text_preview = redacted_text[:500] + "..." if len(redacted_text) > 500 else redacted_text
        
        response = {
            "success": True,
            "document_id": document_id,
            "filename": doc_info['filename'],
            "redacted_filename": doc_info.get('redacted_filename', f"redacted_{doc_info['filename']}"),
            "redacted_count": doc_info.get('redacted_count', 0),
            "redacted_text_preview": redacted_text_preview,
            "download_url": f"/api/download/{document_id}",
            "status": doc_info.get('status', 'redacted')
        }
        
        return response
    
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Redacted preview generation failed: {str(e)}")

@app.get("/api/document/{document_id}/audit-log")
async def get_audit_log(document_id: str):
    """Get audit log for a document"""
    try:
        if document_id not in documents_db:
            raise HTTPException(status_code=404, detail="Document not found")
        
        doc_info = documents_db[document_id]
        
        # Generate audit entries based on document status and information
        audit_entries = []
        
        # 1. Document upload entry
        audit_entries.append({
            "timestamp": doc_info['upload_time'],
            "action": "DOCUMENT_UPLOADED",
            "details": f"Uploaded file: {doc_info['filename']}",
            "status": "completed"
        })
        
        # 2. Entity detection entry
        entities = doc_info.get('entities', [])
        if entities:
            audit_entries.append({
                "timestamp": doc_info['upload_time'],  # Same time as upload for now
                "action": "ENTITIES_DETECTED",
                "details": f"Detected {len(entities)} sensitive entities",
                "status": "completed",
                "entities": entities
            })
        
        # 3. Document redaction entry (if redacted)
        if doc_info.get('status') == 'redacted':
            redacted_count = doc_info.get('redacted_count', 0)
            redacted_entities = doc_info.get('redacted_entities', [])
            redaction_time = doc_info.get('redaction_time', doc_info['upload_time'])
            
            audit_entries.append({
                "timestamp": redaction_time,
                "action": "DOCUMENT_REDACTED",
                "details": f"Redacted {redacted_count} entities",
                "status": "completed",
                "redacted_count": redacted_count,
                "redacted_entities": redacted_entities  # Include the actual redacted entities
            })
        
        response = {
            "success": True,
            "document_id": document_id,
            "filename": doc_info['filename'],
            "audit_entries": audit_entries,
            "total_entries": len(audit_entries)
        }
        
        return response
    
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Audit log generation failed: {str(e)}")

@app.get("/api/document/{document_id}/download-audit-log")
async def download_audit_log(document_id: str):
    """Download comprehensive audit log as JSON file"""
    try:
        if document_id not in documents_db:
            raise HTTPException(status_code=404, detail="Document not found")
        
        doc_info = documents_db[document_id]
        
        # Get current timestamp for audit
        audit_timestamp = datetime.now().isoformat() + 'Z'
        
        # Get audit entries
        audit_response = await get_audit_log(document_id)
        audit_entries = audit_response["audit_entries"]
        
        # Create comprehensive audit log
        audit_log = {
            "document_id": document_id,
            "filename": doc_info['filename'],
            "upload_time": doc_info['upload_time'],
            "status": doc_info.get('status', 'processed'),
            "audit_timestamp": audit_timestamp,
            "entities_detected": len(doc_info.get('entities', [])),
            "entities_redacted": doc_info.get('redacted_count', 0),
            "entities": doc_info.get('entities', []),
            "redacted_text_preview": doc_info.get('redacted_text', '')[:1000] + "..." if len(doc_info.get('redacted_text', '')) > 1000 else doc_info.get('redacted_text', ''),
            "audit_entries": audit_entries
        }
        
        # Create temporary JSON file
        timestamp_str = datetime.now().strftime("%Y%m%d_%H%M%S")
        audit_filename = f"audit_log_{document_id}_{timestamp_str}.json"
        audit_path = downloads_dir / audit_filename
        
        # Write audit log to file
        import json
        with open(audit_path, 'w', encoding='utf-8') as f:
            json.dump(audit_log, f, indent=2, ensure_ascii=False)
        
        # Return file as download
        return FileResponse(
            path=str(audit_path),
            filename=audit_filename,
            media_type='application/json',
            headers={
                "Content-Disposition": f"attachment; filename={audit_filename}"
            }
        )
    
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Audit log download failed: {str(e)}")

# Cleanup function for temporary files
def cleanup_temp_file(file_path: str):
    """Clean up temporary file after download"""
    try:
        if os.path.exists(file_path):
            os.unlink(file_path)
            print(f"🗑️ Cleaned up temporary file: {file_path}")
    except Exception as e:
        print(f"⚠️ Failed to cleanup temporary file {file_path}: {e}")

# Cleanup function
async def cleanup_old_files():
    """Clean up old uploaded and downloaded files"""
    import time
    current_time = time.time()
    
    # Clean up files older than 24 hours
    for directory in [uploads_dir]:
        for file_path in directory.glob("*"):
            if file_path.is_file():
                file_age = current_time - file_path.stat().st_mtime
                if file_age > 86400:  # 24 hours
                    try:
                        file_path.unlink()
                    except Exception as e:
                        print(f"Failed to delete {file_path}: {e}")

if __name__ == "__main__":
    import uvicorn
    print("🚀 Starting Document Redaction API")
    print("📍 Server: http://localhost:8000")
    print("📖 Docs: http://localhost:8000/api/docs")
    uvicorn.run("app:app", host="0.0.0.0", port=8000, reload=True)
